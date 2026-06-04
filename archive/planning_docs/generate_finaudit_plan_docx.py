from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = "FinAudit-Graph_7天开发执行计划书.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, widths_cm):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            cell = row.cells[idx]
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 567)))
            tc_w.set(qn("w:type"), "dxa")


def style_table(table, header_fill="1F4E79", body_fill=None):
    table.style = "Table Grid"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_fill)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
    set_repeat_table_header(table.rows[0])
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            if row_idx > 0 and body_fill and row_idx % 2 == 0:
                set_cell_shading(cell, body_fill)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        if level == 1:
            run.font.color.rgb = RGBColor(31, 78, 121)
        elif level == 2:
            run.font.color.rgb = RGBColor(79, 129, 189)
    return p


def add_body(doc, text, bold_label=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)
    run.bold = bold_label
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10)
    return p


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def apply_document_styles(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    add_page_number(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    title = p.add_run("FinAudit-Graph\n7 天开发执行计划书")
    title.bold = True
    title.font.name = "Arial"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    title.font.size = Pt(28)
    title.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(18)
    r = subtitle.add_run("基于大模型智能体的端到端财务审计与合规审查系统")
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(89, 89, 89)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    widths = [4.0, 9.0]
    labels = ["项目定位", "执行周期", "核心技术栈", "计划用途"]
    values = [
        "冲刺型毕业设计开发与答辩落地",
        "7 天",
        "Label Studio、LLaMA Factory、LangChain 1.x、LangGraph、Neo4j、RAG、N8N、飞书多维表格、Streamlit",
        "指导开发、联调、演示录制、答辩复盘与评分点呈现",
    ]
    for i, (label, value) in enumerate(zip(labels, values)):
        set_cell_text(meta.cell(i, 0), label, bold=True, color="1F4E79", size=10)
        set_cell_text(meta.cell(i, 1), value, size=10)
        set_cell_shading(meta.cell(i, 0), "D9EAF7")
    set_table_width(meta, widths)
    style_table(meta, header_fill="D9EAF7")

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run("版本：V1.0    输出文件：FinAudit-Graph_7天开发执行计划书.docx")
    nr.font.name = "Arial"
    nr._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    nr.font.size = Pt(9)
    nr.font.color.rgb = RGBColor(100, 100, 100)
    note.add_run().add_break(WD_BREAK.PAGE)


def add_overview(doc):
    add_heading(doc, "一、执行目标与成功标准", 1)
    add_body(
        doc,
        "本计划将原《智能审计规划》中的架构设想转化为 7 天可执行开发路径，目标是在毕业设计冲刺周期内完成一个可演示、可解释、可答辩的 FinAudit-Graph 原型系统。",
    )
    add_bullet(doc, "系统可完成“上传材料 -> 智能体分析 -> 图谱/RAG 增强 -> 风险输出 -> 自动化记录 -> 报告生成”的端到端闭环演示。")
    add_bullet(doc, "至少形成一个 LangGraph 四节点核心应用、一个 Neo4j 企业关联样本图谱、一个本地审计准则 RAG 演示库。")
    add_bullet(doc, "完成 Streamlit 前端、N8N Webhook、飞书多维表格写入和报告生成链路的可运行演示。")
    add_bullet(doc, "准备 5 到 8 分钟演示视频、10 到 15 页 PPT 和答辩讲解材料，突出技术路线、业务价值和工程完整性。")


def add_schedule(doc):
    add_heading(doc, "二、7 天开发排期总表", 1)
    rows = [
        ["Day 1", "方案冻结与工程环境", "建立仓库、虚拟环境、依赖清单、目录结构；画出技术路线图；确定演示数据格式。", "Git 仓库、README、环境说明、架构图草稿", "本地可启动基础工程；答辩中可解释总体架构", "依赖安装失败：优先保留最小 LangGraph 可运行骨架"],
        ["Day 2", "数据标注与微调演示", "生成或整理 200 条审计问答/风险分类样本；导入 Label Studio；准备 LLaMA Factory LoRA 演示配置。", "标注样本、Label Studio 截图、训练配置", "能展示标注流程和微调前后对比思路", "训练时间不足：用小样本和日志截图证明流程"],
        ["Day 3", "LangGraph 核心四节点", "实现 Data_Parser、Graph_Searcher、Compliance_Checker、Report_Generator；完成状态流转和 invoke 测试。", "核心 Python 代码、模拟 PDF 输入、Markdown 报告输出", "四节点顺序执行且最终报告可打印", "真实解析复杂：先用模拟抽取保证主链路可演示"],
        ["Day 4", "Neo4j 与 RAG 增强", "导入企业股权样本；编写 Cypher 查询；建立审计准则向量库；把图谱和法规检索结果注入风险判断。", "Neo4j 样本图、Cypher 查询、RAG 检索脚本", "能演示关联方穿透和准则命中", "数据库联动不稳：保留本地 JSON fallback"],
        ["Day 5", "前端与自动化闭环", "开发 Streamlit 上传页；触发后端分析；N8N Webhook 接收结果；写入飞书多维表格。", "前端页面、N8N 工作流、飞书表格记录", "上传一次材料后生成一条审计疑点记录", "飞书权限问题：准备本地表格/截图备用演示"],
        ["Day 6", "报告生成与联调", "整合风险点、关联方和建议结论；生成 Word/PDF 报告；做端到端压力联调和异常日志修复。", "10 页报告样稿、联调记录、异常处理说明", "完整链路连续运行 3 次无阻断", "报告页数不足：补充风险解释、依据和整改建议"],
        ["Day 7", "答辩冲刺与复盘", "录制 5 到 8 分钟演示视频；制作 PPT；准备问题清单；提炼创新点、局限和后续优化。", "演示视频、PPT、答辩稿、FAQ", "答辩材料完整且演示节奏顺畅", "现场环境不稳：准备录屏和静态截图兜底"],
    ]
    headers = ["日期", "阶段目标", "具体任务", "当日产出", "验收标准", "风险与补救"]
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True, color="FFFFFF", size=8.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8)
    set_table_width(table, [1.2, 2.2, 3.7, 2.7, 2.7, 3.1])
    style_table(table, body_fill="F7FBFE")


def add_modules(doc):
    add_heading(doc, "三、核心模块开发清单", 1)
    rows = [
        ["Data_Parser", "接收 PDF/合同/财报路径，完成文本抽取、企业名称识别和关键财务指标清洗。", "raw_document_path、文件元信息", "parsed_financial_data", "先实现模拟抽取，再逐步接入 pdfplumber、pymupdf 或 OCR。"],
        ["Graph_Searcher", "连接 Neo4j，执行关联方和股权穿透查询，识别未披露控制关系。", "企业名称、股东/交易方线索", "discovered_related_parties", "保留标准 Cypher 查询模板，并提供样本图谱数据。"],
        ["Compliance_Checker", "结合 RAG 审计准则和微调模型，对财务指标、关联方和历史案例进行风险分类。", "parsed_financial_data、关联方、准则片段", "audit_risks_found", "用结构化 JSON 输出风险等级、依据、建议。"],
        ["Report_Generator", "汇总风险点、法规依据和整改建议，生成 Markdown/Word 审计综述。", "全局 AuditSystemState", "final_audit_summary", "报告要适合答辩展示，包含业务解释和技术依据。"],
        ["自动化工作流", "N8N 接收 Webhook，调用后端分析，将风险记录写入飞书多维表格。", "前端上传事件、审计结果 JSON", "飞书待跟进审计疑点记录", "先打通 HTTP 节点，再补鉴权和异常重试。"],
        ["前端演示页", "Streamlit 提供上传、分析进度、风险结果、报告下载入口。", "用户上传材料、按钮操作", "可交互演示界面", "界面优先简洁稳定，突出端到端闭环。"],
    ]
    headers = ["模块", "职责", "输入", "输出", "实现要点"]
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True, color="FFFFFF", size=9)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8.5)
    set_table_width(table, [2.4, 3.9, 2.8, 2.8, 3.8])
    style_table(table, body_fill="F7FBFE")


def add_deliverables(doc):
    add_heading(doc, "四、交付物矩阵与验收清单", 1)
    rows = [
        ["工程代码", "LangGraph 四节点核心应用、Streamlit 前端、Neo4j/RAG 接入脚本", "能本地运行；README 可指导复现；关键函数有中文注释"],
        ["数据与模型", "标注样本、Label Studio 演示截图、LLaMA Factory 配置或训练日志", "能说明训练数据来源、标签含义和微调价值"],
        ["自动化链路", "N8N Webhook 工作流、飞书多维表格风险记录、告警或报告生成节点", "上传一次样例后能产生结构化记录"],
        ["审计报告", "Markdown/Word/PDF 风险审计报告样稿", "包含风险等级、涉及主体、依据、建议和综合结论"],
        ["答辩材料", "10 到 15 页 PPT、5 到 8 分钟演示视频、FAQ 问题清单", "覆盖背景、架构、核心代码、演示、创新和局限"],
    ]
    table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["类别", "内容", "验收标准"]):
        set_cell_text(table.cell(0, i), h, bold=True, color="FFFFFF", size=9)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=9)
    set_table_width(table, [2.7, 6.5, 6.8])
    style_table(table, body_fill="F7FBFE")

    add_heading(doc, "五、每日执行检查点", 1)
    checks = [
        "每天结束前提交一次代码或材料快照，并记录当天已完成、未完成和阻塞项。",
        "所有演示链路都要准备截图或录屏兜底，避免答辩现场网络、权限或服务启动失败。",
        "核心代码必须保留中文工程注释，尤其是 LangGraph 状态定义、节点职责、Cypher 查询和 RAG 调用位置。",
        "所有输出结果尽量结构化为 JSON 或 Markdown，方便写入飞书和生成审计报告。",
        "答辩展示优先讲清楚业务问题、技术方案、运行链路和风险识别结果，不把时间耗在安装或调试细节上。",
    ]
    for item in checks:
        add_bullet(doc, item)


def add_risks(doc):
    add_heading(doc, "六、风险清单与应对策略", 1)
    rows = [
        ["依赖安装或版本冲突", "LangChain/LangGraph、Neo4j Driver、LLaMA Factory 版本不兼容", "中", "固定 requirements；先做最小可运行版本；关键依赖记录截图"],
        ["微调训练时间不足", "7 天周期内无法完成完整训练或评估", "高", "用小样本 LoRA 演示流程；保留训练日志；用 API 模型完成主链路"],
        ["Neo4j/RAG 联动不稳定", "数据库、向量库或检索链路影响演示连续性", "中", "准备样本 JSON fallback；答辩时展示图谱截图和查询结果"],
        ["N8N/飞书权限问题", "Webhook、Token、表格权限导致无法写入", "中", "提前创建测试表；保留本地日志和飞书截图；必要时用模拟 HTTP 节点"],
        ["报告生成质量不足", "风险结论过短、页数不足或业务解释不够", "中", "固定报告模板；加入法规依据、整改建议和复核步骤"],
        ["答辩现场演示失败", "服务启动、网络、投屏或接口调用失败", "高", "准备录屏、截图和本地 mock；现场演示只跑最稳定路径"],
    ]
    table = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["风险项", "表现", "优先级", "应对策略"]):
        set_cell_text(table.cell(0, i), h, bold=True, color="FFFFFF", size=9)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8.5)
    set_table_width(table, [3.6, 4.8, 1.6, 6.0])
    style_table(table, body_fill="F7FBFE")


def add_scoring(doc):
    add_heading(doc, "七、毕设评分导向与答辩呈现", 1)
    add_body(doc, "答辩时建议围绕“业务痛点明确、技术链路完整、智能体编排清晰、演示闭环可运行”四个维度组织讲解。")
    rows = [
        ["技术创新点", "用 LangGraph 显式编排多智能体，结合 RAG 与 Neo4j 做审计风险增强判断。"],
        ["工程完整性", "从数据标注、模型微调、后端分析、前端交互到自动化记录形成端到端闭环。"],
        ["业务贴合度", "聚焦关联交易、财务异常、内控缺陷、法规依据和整改建议，贴近真实审计流程。"],
        ["可解释性", "每个风险点输出依据、关联方、风险等级和建议，便于答辩说明模型不是黑箱。"],
        ["演示亮点", "上传样例财报后，系统自动完成分析、写入飞书并生成报告，形成视觉化闭环。"],
    ]
    table = doc.add_table(rows=1, cols=2)
    set_cell_text(table.cell(0, 0), "评分维度", bold=True, color="FFFFFF", size=9)
    set_cell_text(table.cell(0, 1), "呈现方式", bold=True, color="FFFFFF", size=9)
    for row in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], row[0], bold=True, color="1F4E79", size=9)
        set_cell_text(cells[1], row[1], size=9)
    set_table_width(table, [3.2, 12.8])
    style_table(table, body_fill="F7FBFE")


def build():
    doc = Document()
    apply_document_styles(doc)
    add_title_page(doc)
    add_overview(doc)
    add_schedule(doc)
    add_modules(doc)
    add_deliverables(doc)
    add_risks(doc)
    add_scoring(doc)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
