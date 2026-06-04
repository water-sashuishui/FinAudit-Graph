from __future__ import annotations

import json
from pathlib import Path

from .automation import build_n8n_payload
from .state import AuditSystemState


def _safe_filename(value: str) -> str:
    cleaned = []
    previous_was_separator = False
    aliases = {
        "华辰智能装备股份有限公司": "huachen_intelligent_equipment",
    }
    value = aliases.get(value, value)
    for character in value.lower():
        if character.isascii() and character.isalnum():
            cleaned.append(character)
            previous_was_separator = False
        elif not previous_was_separator:
            cleaned.append("_")
            previous_was_separator = True
    slug = "".join(cleaned).strip("_")
    return slug or "audit_subject"


def build_full_report_markdown(state: AuditSystemState) -> str:
    """Build a formal MVP audit report from the workflow state."""
    parsed = state.get("parsed_financial_data", {})
    related_parties = state.get("discovered_related_parties", [])
    risks = state.get("audit_risks_found", [])
    high_risk_count = sum(1 for risk in risks if risk.get("severity") == "高")
    company_name = parsed.get("company_name", "未知企业")
    payload = build_n8n_payload(state)

    risk_blocks = []
    for index, risk in enumerate(risks, start=1):
        risk_blocks.append(
            "\n".join(
                [
                    f"### {index}. {risk['risk_type']}（{risk['severity']}风险）",
                    f"- 风险依据：{risk['evidence']}",
                    f"- 审计准则/关注点：{risk['audit_basis']}",
                    f"- 整改建议：{risk['recommendation']}",
                ]
            )
        )

    related_party_lines = [
        f"- {item['name']}：{item['relation']}，穿透深度 {item['depth']}。证据：{item['evidence']}"
        for item in related_parties
    ] or ["- 暂未发现可疑关联方。"]

    key_clues = parsed.get("key_clues", [])
    clue_lines = [f"- {item}" for item in key_clues] or ["- 暂无关键线索。"]

    payload_summary = {
        "company_name": payload["company_name"],
        "reporting_year": payload["reporting_year"],
        "risk_count": payload["risk_count"],
        "high_risk_count": payload["high_risk_count"],
        "automation_mode": "dry-run 或真实 N8N webhook",
    }

    return f"""# 企业合规风控审计报告

## 一、执行摘要

本报告由 FinAudit-Graph MVP 自动生成。系统围绕 {company_name} 的财务指标、潜在关联方和审计准则依据进行多智能体分析，共识别 {len(risks)} 个审计风险点，其中高风险 {high_risk_count} 个。当前结论适合作为答辩演示、审计底稿初筛和后续人工复核线索。

## 二、被审计对象与输入材料

- 企业名称：{company_name}
- 报告年度：{parsed.get("reporting_year", "未知")}
- 原始文件：{parsed.get("source_file", "未提供")}
- 收入增长率：{parsed.get("revenue_growth_rate", "N/A")}%
- 经营现金流增长率：{parsed.get("operating_cashflow_growth_rate", "N/A")}%
- 毛利率：{parsed.get("gross_margin_rate", "N/A")}%
- 应收账款增长率：{parsed.get("accounts_receivable_growth_rate", "N/A")}%

## 三、核心财务线索

{chr(10).join(clue_lines)}

## 四、关联方穿透结果

{chr(10).join(related_party_lines)}

## 五、整改建议与复核计划

1. 对收入确认执行截止性测试，抽查期末大额合同、验收单、发票和回款记录。
2. 对潜在关联方执行股权穿透、资金流水复核和高管交叉任职核查。
3. 对毛利率异常和成本结转执行分析性复核，补充检查成本计算单和出入库记录。
4. 对内控缺陷线索开展流程穿行测试，确认审批、验收、对账和复核证据是否完整。

## 六、审计风险明细

{chr(10).join(risk_blocks)}

## 七、自动化记录 Payload 摘要

```json
{json.dumps(payload_summary, ensure_ascii=False, indent=2)}
```

## 八、综合结论

系统建议将 {company_name} 列为高优先级复核对象。当前 MVP 已经形成“材料输入、图谱线索、向量数据库 RAG 依据、风险判断、自动化 payload、报告输出”的端到端闭环；真实生产环境中仍需接入完整 PDF 解析、真实 Neo4j、生产级向量数据库、N8N 和飞书鉴权。
"""


def save_markdown_report(state: AuditSystemState, output_dir: str | Path = "outputs") -> Path:
    """Save the final audit summary as a Markdown report."""
    target_dir = Path(output_dir)
    target_dir.mkdir(parents=True, exist_ok=True)
    parsed = state.get("parsed_financial_data", {})
    company_name = parsed.get("company_name", "unknown_company")
    safe_name = _safe_filename(company_name)
    if safe_name == "audit_subject":
        safe_name = _safe_filename(Path(parsed.get("source_file", "audit_subject")).stem)
    path = target_dir / f"{safe_name}_audit_report.md"
    path.write_text(build_full_report_markdown(state), encoding="utf-8")
    return path


def save_docx_report(state: AuditSystemState, output_dir: str | Path = "outputs") -> Path | None:
    """Save a formal Word report if python-docx is available."""
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except ImportError:
        return None

    target_dir = Path(output_dir)
    target_dir.mkdir(parents=True, exist_ok=True)
    parsed = state.get("parsed_financial_data", {})
    company_name = parsed.get("company_name", "unknown_company")
    safe_name = _safe_filename(company_name)
    if safe_name == "audit_subject":
        safe_name = _safe_filename(Path(parsed.get("source_file", "audit_subject")).stem)
    path = target_dir / f"{safe_name}_audit_report.docx"

    def set_font(paragraph, east_asia: str = "微软雅黑") -> None:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east_asia)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)

    title = doc.add_heading("企业合规风控审计报告", level=0)
    set_font(title)
    subtitle = doc.add_paragraph("FinAudit-Graph MVP 自动生成")
    set_font(subtitle)

    markdown = build_full_report_markdown(state)
    for line in markdown.splitlines():
        if not line.strip() or line.startswith("```"):
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            paragraph = doc.add_heading(line.replace("## ", "", 1), level=1)
        elif line.startswith("### "):
            paragraph = doc.add_heading(line.replace("### ", "", 1), level=2)
        elif line.startswith("- ") or line[0:2] in {"1.", "2.", "3.", "4."}:
            paragraph = doc.add_paragraph(line, style="List Bullet")
        else:
            paragraph = doc.add_paragraph(line)
        set_font(paragraph)

    try:
        doc.save(path)
    except OSError:
        return None
    return path


def save_reports(state: AuditSystemState, output_dir: str | Path = "outputs") -> dict[str, str | None]:
    markdown_path = save_markdown_report(state, output_dir)
    docx_path = save_docx_report(state, output_dir)
    return {
        "markdown": str(markdown_path),
        "docx": str(docx_path) if docx_path else None,
    }
